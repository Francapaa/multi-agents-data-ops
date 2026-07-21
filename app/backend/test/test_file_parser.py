import io
import os
import pytest

from services.file_parser import (
    parse_file,
    _parse_docx,
    SUPPORTED_EXTENSIONS,
    MAX_FILE_SIZE,
    UnsupportedFormatError,
    FileTooLargeError,
)


class TestConstants:
    def test_supported_extensions(self):
        assert ".txt" in SUPPORTED_EXTENSIONS
        assert ".pdf" in SUPPORTED_EXTENSIONS
        assert ".docx" in SUPPORTED_EXTENSIONS

    def test_max_file_size(self):
        assert MAX_FILE_SIZE == 10 * 1024 * 1024


class TestParseFile:
    def test_txt_file_parsed_correctly(self):
        content = "Hello world\nThis is a test."
        result = parse_file(content.encode("utf-8"), "test.txt")
        assert result == content

    def test_unsupported_format_raises_error(self):
        with pytest.raises(UnsupportedFormatError):
            parse_file(b"content", "test.csv")

    def test_file_too_large_raises_error(self):
        large_bytes = b"x" * (MAX_FILE_SIZE + 1)
        with pytest.raises(FileTooLargeError):
            parse_file(large_bytes, "test.txt")

    def test_file_at_max_size_boundary(self):
        bytes_at_limit = b"x" * MAX_FILE_SIZE
        result = parse_file(bytes_at_limit, "test.txt")
        assert len(result) == MAX_FILE_SIZE

    def test_empty_txt_file(self):
        result = parse_file(b"", "test.txt")
        assert result == ""

    def test_unicode_text_in_txt(self):
        content = "Hola mundo! \u00a1Ol\u00e9! \u4e2d\u6587"
        result = parse_file(content.encode("utf-8"), "test.txt")
        assert result == content


class TestParseDocx:
    def test_docx_with_single_paragraph(self):
        from docx import Document

        doc = Document()
        doc.add_paragraph("Hello world")
        buf = io.BytesIO()
        doc.save(buf)
        result = _parse_docx(buf.getvalue())
        assert "Hello world" in result

    def test_docx_with_multiple_paragraphs(self):
        from docx import Document

        doc = Document()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph")
        buf = io.BytesIO()
        doc.save(buf)
        result = _parse_docx(buf.getvalue())
        assert "First paragraph" in result
        assert "Second paragraph" in result

    def test_empty_docx(self):
        from docx import Document

        doc = Document()
        buf = io.BytesIO()
        doc.save(buf)
        result = _parse_docx(buf.getvalue())
        assert result == ""
